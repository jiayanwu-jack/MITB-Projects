from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("Project Portfolio Report.docx")
NAVY = "17365D"
BLUE = "2E75B6"
PALE = "EAF1F8"
LIGHT = "F2F4F7"
GRAY = "666666"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc.get_or_add_tcPr()
    tc_mar = tc.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, color=None, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Aptos"
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    margins(cell)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.style = "Table Grid"
    for i, (header, width) in enumerate(zip(headers, widths)):
        table.columns[i].width = Inches(width)
        set_cell_text(table.rows[0].cells[i], header, bold=True, color="FFFFFF")
        shade(table.rows[0].cells[i], NAVY)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for i, (value, width) in enumerate(zip(values, widths)):
            cells[i].width = Inches(width)
            set_cell_text(cells[i], value)
            if row_index % 2:
                shade(cells[i], LIGHT)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_project(doc, number, title, subtitle, objective, approach, findings, outputs, caveat):
    doc.add_heading(f"{number}. {title}", level=1)
    p = doc.add_paragraph()
    p.style = doc.styles["Subtitle"]
    p.add_run(subtitle)
    doc.add_heading("Objective", level=2)
    doc.add_paragraph(objective)
    doc.add_heading("Approach", level=2)
    for item in approach:
        add_bullet(doc, item)
    doc.add_heading("Key findings", level=2)
    for item in findings:
        add_bullet(doc, item)
    doc.add_heading("Deliverables and skills demonstrated", level=2)
    doc.add_paragraph(outputs)
    p = doc.add_paragraph()
    r = p.add_run("Interpretation note: ")
    r.bold = True
    p.add_run(caveat)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.08
for name, size, color, before, after in (
    ("Title", 25, NAVY, 0, 8),
    ("Subtitle", 11.5, GRAY, 0, 12),
    ("Heading 1", 16, BLUE, 14, 6),
    ("Heading 2", 12, NAVY, 9, 3),
):
    style = styles[name]
    style.font.name = "Aptos Display" if name in ("Title", "Heading 1") else "Aptos"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
styles["Heading 1"].font.bold = True
styles["Heading 2"].font.bold = True
styles["List Bullet"].font.name = "Aptos"
styles["List Bullet"].font.size = Pt(10.5)
styles["List Bullet"].paragraph_format.left_indent = Inches(0.28)
styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.16)
styles["List Bullet"].paragraph_format.space_after = Pt(3)

header = section.header.paragraphs[0]
header.text = "MITB PROJECT PORTFOLIO"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header.runs[0].font.name = "Aptos"
header.runs[0].font.size = Pt(8)
header.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("Prepared from the contents of the Projects folder | 11 June 2026")
run.font.name = "Aptos"
run.font.size = Pt(8)
run.font.color.rgb = RGBColor.from_string(GRAY)

doc.add_paragraph("PROJECT PORTFOLIO REPORT", style="Title")
doc.add_paragraph(
    "Master of IT in Business project work across analytics, machine learning, "
    "financial services, database optimization, decision modelling, and statistics",
    style="Subtitle",
)
doc.add_paragraph("Prepared for: Wu Jiayan")
doc.add_paragraph("Source scope: Six subfolders and eleven project artifacts in the local Projects directory")

callout = doc.add_table(rows=1, cols=1)
callout.autofit = False
callout.columns[0].width = Inches(6.55)
cell = callout.cell(0, 0)
set_cell_text(
    cell,
    "Portfolio theme: translating data into operational decisions through forecasting, "
    "optimization, statistical inference, system performance engineering, and responsible AI.",
    bold=True,
    color=NAVY,
    size=11,
)
shade(cell, PALE)

doc.add_heading("Executive Summary", level=1)
doc.add_paragraph(
    "The Projects folder documents six applied academic projects. Together, they show a "
    "consistent progression from problem framing and data preparation to model construction, "
    "evaluation, and decision-oriented recommendations. The work covers Python notebooks, "
    "Excel forecasting and Solver models, statistical analysis, SQL indexing benchmarks, "
    "AI-assisted compliance surveillance, and presentation of business implications."
)
add_table(
    doc,
    ["Project area", "Decision problem", "Primary evidence", "Headline result"],
    [
        ("Applied Machine Learning", "Forecast future HDB resale prices", "DOCX + Jupyter notebook", "Stacking ensemble: R² 0.7314; RMSE 865.13"),
        ("Data Science for Business", "Optimize Airbnb nightly prices", "PPTX + script + notebook", "Random Forest ROC-AUC 0.8441; case-level revenue uplift"),
        ("Financial Services", "Detect misconduct in trader chats", "Individual lab report", "AI intent analysis outperformed static keyword review"),
        ("Query Optimization", "Accelerate text search", "Benchmark presentation", "Average runtime reduction about 90%; up to 21x faster"),
        ("Spreadsheet Modelling", "Test Singapore water sufficiency to 2060", "XLSX + report + slides", "Supply remains sufficient in model; tariff path protects 10% buffer"),
        ("Statistical Thinking", "Assess healthcare adequacy for ageing", "Group report", "Capacity broadly tracks ageing; psychiatric beds lag"),
    ],
    [1.35, 1.75, 1.45, 2.0],
)

doc.add_heading("Portfolio-Wide Strengths", level=1)
for item in (
    "Decision orientation: each project converts analysis into a recommendation, policy implication, or deployable workflow.",
    "Methodological breadth: supervised learning, ensemble methods, time-series validation, optimization, sensitivity analysis, inference, and database benchmarking.",
    "Practical tooling: Python, pandas, scikit-learn, Jupyter, Excel Solver and Data Tables, SQL full-text indexing, and Google AI Studio.",
    "Communication: technical findings are consistently translated into presentations, reports, demonstrations, and quantified business outcomes.",
):
    add_bullet(doc, item)

add_project(
    doc,
    1,
    "Applied Machine Learning",
    "Building the best HDB resale price estimator using ensemble learning",
    "Determine whether ensemble methods outperform standalone baselines when predicting genuinely future Singapore HDB resale transactions.",
    [
        "Used a chronological split to limit temporal leakage: training through 2023 (169,150 rows), validation on 2024 (27,832), and testing on 2025 onward (27,683).",
        "Engineered property, location, lease, time, and structural-event features, including MRT and school distances and post-COVID/cooling-measure indicators.",
        "Tuned linear, tree, random forest, and gradient boosting models with five-fold time-series cross-validation.",
        "Built a stacking regressor using tuned Random Forest, Gradient Boosting, and Ridge components with a Ridge meta-model.",
    ],
    [
        "The optimized stacking ensemble achieved the best future-test performance: MAE 725.87, RMSE 865.13, and R² 0.7314.",
        "Against Ridge, the ensemble reduced RMSE by about 155 per sqm and MAE by about 124 per sqm, while improving R² by roughly 10 percentage points.",
        "Gradient Boosting was the strongest individual model (R² 0.6858), confirming the value of nonlinear relationships before stacking.",
    ],
    "A research report and executable notebook demonstrate leakage-aware validation, feature engineering, hyperparameter tuning, ensemble design, and comparative model evaluation.",
    "Reported accuracy applies to the specific temporal holdout and feature set. Deployment would require drift monitoring, periodic retraining, and confirmation that all input variables are available at prediction time.",
)

add_project(
    doc,
    2,
    "Data Science for Business",
    "Cape Town Airbnb dynamic pricing and expected-revenue optimization",
    "Build a pricing engine that recommends a nightly rate by balancing booking probability against price, rather than maximizing occupancy alone.",
    [
        "Merged listing and calendar data and created market-median anchor prices by neighbourhood, room type, and bedroom count.",
        "Trained a Random Forest booking classifier on a one-million-row sample with structural, location, calendar, host, review, and relative-price features.",
        "Evaluated candidate price points and selected the price maximizing expected revenue: price multiplied by predicted booking probability.",
        "Applied the model to 22 new listings and demonstrated date-specific recommendations for listing 231402.",
    ],
    [
        "The booking model achieved ROC-AUC 0.8441 on the notebook's holdout sample.",
        "For 24 June 2026, listing 231402 moved from a R1,200 anchor to a R2,400 recommendation, with 24.58% booking probability and expected daily revenue of R589.82.",
        "The presentation reports a December recommendation of R1,900 with stronger booking probability and expected revenue of R767, illustrating date-level demand effects.",
        "The project connects host revenue optimization to platform commission upside and targeted Smart Pricing nudges.",
    ],
    "The notebook, presentation, and speaking script form an end-to-end business analytics package: predictive modelling, optimization, batch scoring, scenario demonstration, and commercial storytelling.",
    "The notebook uses a random train-test split and market medians as proxies for new listings. Production use would benefit from time-based validation, calibration checks, guardrails on extreme price multiples, and controlled experimentation.",
)

add_project(
    doc,
    3,
    "Data Science for Financial Services",
    "AI-assisted trade surveillance for misconduct in communication logs",
    "Compare manual keyword searching with an AI-assisted workflow for detecting market manipulation, front-running, confidential-information leakage, and coded communications.",
    [
        "Used Excel keyword searches as a baseline and manually reviewed surrounding conversational context.",
        "Built a TradeGuard AI app in Google AI Studio with a categorized surveillance dictionary, full-scan workflow, and segment-level reasoning.",
        "Assessed contextual intent, linked multi-step behavior, discovered possible new codewords, and mapped suspicious behavior to regulatory concepts.",
    ],
    [
        "Manual search was transparent but labor-intensive and prone to false alarms because it could not interpret context.",
        "AI analysis distinguished legitimate market language from suspicious intent and linked related actions, such as a confidentiality breach followed by front-running.",
        "The work argues for dynamic dictionaries and human review because trader language and concealment strategies evolve.",
    ],
    "The report demonstrates regulatory problem framing, qualitative comparison of surveillance methods, prompt/app design, contextual language analysis, and compliance-focused communication.",
    "The report provides illustrative findings rather than a labelled benchmark with precision, recall, or false-positive rates. An operational system would require validated test sets, audit logs, privacy controls, explainability, and mandatory human adjudication.",
)

add_project(
    doc,
    4,
    "Query Processing and Optimisation",
    "MySQL FULLTEXT indexing for Airbnb listing search",
    "Measure whether full-text indexing improves keyword and phrase search over multi-column wildcard LIKE queries on text-heavy Airbnb data.",
    [
        "Loaded an Airbnb dataset of roughly 36,000 listings and 85 attributes, focusing on natural-language fields.",
        "Compared full scans using LIKE '%term%' with MATCH ... AGAINST queries backed by an inverted full-text index.",
        "Benchmarked 34 realistic search terms and phrases, running each query ten times and comparing runtime and rows examined.",
        "Examined single- versus multi-column indexes, storage growth, write costs, stop-word handling, and result-set selectivity.",
    ],
    [
        "The main presentation reports average optimized runtime of 16.69 ms versus 169.58 ms, an average improvement of about 90%.",
        "FULLTEXT examined approximately one indexed entry versus 33,407 rows for non-optimized searches in the benchmark summary.",
        "Performance gains varied by selectivity: examples ranged from about 3x for 'washer dryer' to about 700x for 'highly recommend'.",
        "The recommended use case is a read-heavy, text-dense platform where search scalability outweighs added storage and slower writes.",
    ],
    "The presentation demonstrates index design, query-plan reasoning, experimental benchmarking, complexity analysis, and a balanced evaluation of operational trade-offs.",
    "FULLTEXT and substring LIKE do not always have identical matching semantics. Benchmark interpretation should therefore verify relevance equivalence, cache conditions, database configuration, and repeatability at larger scales.",
)

add_project(
    doc,
    5,
    "Spreadsheet Modelling for Decision Making",
    "Singapore water sufficiency, tariff optimization, and sensitivity analysis to 2060",
    "Evaluate whether planned water supply can support population and economic growth through 2060, and identify tariff adjustments that preserve an adequate reserve margin.",
    [
        "Forecast population and per-capita consumption with uncertainty bounds and model demand across domestic and industrial use.",
        "Represent the Four National Taps: local catchment, imported water, NEWater, and desalination, including rainfall variability, reliability assumptions, capacity factors, and planned plants.",
        "Use price elasticity (-0.2 baseline) to connect tariff increases with consumption changes.",
        "Apply Excel Solver to minimize average annual tariff increases while enforcing a 10% reserve margin, then test elasticity from -0.1 to -1.0 with Data Tables.",
    ],
    [
        "The written report concludes that supply remains sufficient through 2060 under current infrastructure plans and moderate growth assumptions.",
        "Modelled supply rises from about 1.03 billion m³ in 2026 to about 1.57 billion m³ from 2035, supported by Tuas Nexus and Kranji-Changi expansions.",
        "In the workbook's high-population/baseline-consumption scenario, cached results show a 2026 reserve margin of 29.0%, rising after new capacity, then converging to the 10% constraint by 2056-2060.",
        "The optimized workbook averages a 3.15% annual tariff increase; later adjustments become steeper, reaching about 36.8% in 2060 under that scenario.",
        "At 2060, weakening elasticity from -0.2 to -0.1 lowers the reserve margin from 10.0% to about 5.8%, while stronger response improves resilience.",
    ],
    "The workbook, report, and slides demonstrate integrated forecasting, scenario design, formula modelling, nonlinear Solver optimization, sensitivity analysis, dashboarding, and policy communication.",
    "The result is scenario-dependent and relies on planned capacity, utilization, import reliability, demand shares, forecast functions, and elasticity assumptions. Tariff affordability and political feasibility should be treated as separate decision criteria.",
)

add_project(
    doc,
    6,
    "Statistical Thinking for Data Science",
    "Assessing whether Singapore healthcare is adequate for an ageing population",
    "Test whether healthcare infrastructure and expenditure have kept pace with growth in Singapore's elderly population, and identify capacity areas that may be lagging.",
    [
        "Combined demographic, expenditure, bed-capacity, occupancy, and admission indicators.",
        "Used descriptive statistics, Pearson correlations and significance tests, simple linear regression, and paired growth-rate t-tests.",
        "Interpreted both system-wide adequacy and segment-specific risks.",
    ],
    [
        "Elderly share was strongly correlated with health expenditure per capita (r=0.959) and acute, community, and nursing-home beds (r=0.976-0.992), with reported p-values below 0.05.",
        "The regression reported R²=0.92, indicating that elderly population share explains 92% of observed variation in healthcare expenditure within the analysed data.",
        "Paired tests found no statistical evidence that elderly growth outpaced acute, community, or nursing-home bed growth.",
        "Psychiatric beds had a negative correlation with elderly share (r=-0.81), highlighting a potential geriatric mental-health capacity gap.",
        "Rising expenditure was also associated with bed occupancy (r=0.81), suggesting that added resources are rapidly absorbed by demand.",
    ],
    "The group report demonstrates hypothesis formulation, inference, regression interpretation, adequacy assessment, and translation of statistical evidence into healthcare planning implications.",
    "Correlation does not establish causality, and trending time-series variables can produce strong shared correlations. Robust follow-up should test stationarity, lags, confounders, per-capita measures, and segment-level service quality.",
)

doc.add_heading("Cross-Project Capability Matrix", level=1)
add_table(
    doc,
    ["Capability", "Evidence across projects"],
    [
        ("Predictive modelling", "HDB regression ensembles; Airbnb booking classification"),
        ("Optimization", "Airbnb expected-revenue search; water-tariff Solver model"),
        ("Validation and testing", "Time-series CV; ROC-AUC; SQL runtime benchmarks; hypothesis tests"),
        ("Uncertainty analysis", "Temporal holdout; forecast bounds; scenarios; elasticity sensitivity"),
        ("Data engineering", "Feature engineering, dataset merging, Office/SQL models, multi-source public data"),
        ("Responsible deployment", "Leakage control, surveillance human review, model caveats, policy constraints"),
        ("Communication", "Reports, slide decks, scripts, dashboards, quantified recommendations"),
    ],
    [1.65, 4.9],
)

doc.add_heading("Overall Assessment and Development Priorities", level=1)
doc.add_paragraph(
    "The portfolio's strongest feature is its applied range: the projects do not stop at model "
    "accuracy, but connect analytical outputs to pricing, infrastructure, compliance, search "
    "performance, housing forecasts, and healthcare planning. The most mature technical work "
    "is the HDB forecasting project because it explicitly guards against temporal leakage and "
    "evaluates a locked future period. The water model stands out for integrating physical "
    "supply, demand behavior, optimization, and sensitivity in a decision interface."
)
for item in (
    "Standardize reproducibility: add README files, data dictionaries, environment specifications, and one-command execution where possible.",
    "Strengthen validation: prefer time-based splits for temporal business problems, calibration for probabilities, and out-of-sample stress tests for policy models.",
    "Add governance artifacts: model cards, assumption registers, privacy controls, fairness checks, and human-review procedures for high-impact use cases.",
    "Quantify uncertainty in recommendations: include confidence intervals, scenario ranges, and sensitivity of business value to model error.",
    "Package selected work into deployable demos or dashboards while retaining the reports as evidence of reasoning and communication.",
):
    add_bullet(doc, item)

doc.add_heading("Appendix: Source Artifacts Reviewed", level=1)
add_table(
    doc,
    ["Subfolder", "Artifacts"],
    [
        ("Applied Machine Learning", "Research Report.docx; RQ4 Ensemble Methods.ipynb"),
        ("Data Science for Business", "Pricing Optimization Model Presentation.pptx; Presentation Script.docx; Pricing Optimization Model.ipynb"),
        ("Data Science for Financial Service", "Individual Lab Report.docx"),
        ("Query Processing and Optimisation", "ISSS625_Team3_Full Text Index.pptx"),
        ("Spreadsheet Modelling for Decision Making", "Water model XLSX; presentation PPTX; final report DOCX"),
        ("Statistical Thinking for Data Science", "IS630 Statistical Thinking Group Project.docx"),
    ],
    [2.15, 4.4],
)
doc.add_paragraph(
    "Scope note: this report summarizes the evidence contained in the listed local artifacts. "
    "It does not independently re-run external datasets, verify citations against current web "
    "sources, or claim results beyond the scenarios and evaluations documented in those files."
)

doc.save(OUT)
print(OUT.resolve())
