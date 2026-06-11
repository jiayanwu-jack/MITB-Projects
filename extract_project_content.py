import json
import re
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from openpyxl import load_workbook


sys.stdout.reconfigure(encoding="utf-8")
ROOT = Path("Projects")


def clean(text):
    return re.sub(r"\s+", " ", str(text)).strip()


for path in ROOT.rglob("*.docx"):
    print(f"\n### DOCX: {path}")
    doc = Document(path)
    for para in doc.paragraphs:
        text = clean(para.text)
        if text:
            print(text)
    for index, table in enumerate(doc.tables, 1):
        print(f"[TABLE {index}]")
        for row in table.rows:
            print(" | ".join(clean(cell.text) for cell in row.cells))


for path in ROOT.rglob("*.pptx"):
    print(f"\n### PPTX: {path}")
    with zipfile.ZipFile(path) as archive:
        slides = sorted(
            (
                name
                for name in archive.namelist()
                if re.fullmatch(r"ppt/slides/slide\d+\.xml", name)
            ),
            key=lambda name: int(re.search(r"(\d+)", Path(name).stem).group(1)),
        )
        for index, name in enumerate(slides, 1):
            root = ET.fromstring(archive.read(name))
            texts = [
                clean(node.text)
                for node in root.iter()
                if node.tag.endswith("}t") and node.text
            ]
            print(f"[SLIDE {index}] {' | '.join(texts)}")


for path in ROOT.rglob("*.ipynb"):
    print(f"\n### NOTEBOOK: {path}")
    notebook = json.loads(path.read_text(encoding="utf-8"))
    for index, cell in enumerate(notebook.get("cells", []), 1):
        source = clean("".join(cell.get("source", [])))
        if source:
            print(f"[{cell.get('cell_type', '').upper()} {index}] {source[:3500]}")
        for output in cell.get("outputs", []):
            text = "".join(output.get("text", []))
            if not text:
                text = "".join(output.get("data", {}).get("text/plain", []))
            text = clean(text)
            if text:
                print(f"[OUTPUT] {text[:2500]}")


for path in ROOT.rglob("*.xlsx"):
    print(f"\n### XLSX: {path}")
    formulas = load_workbook(path, data_only=False, read_only=False)
    values = load_workbook(path, data_only=True, read_only=False)
    print("[SHEETS] " + " | ".join(formulas.sheetnames))
    for name in formulas.sheetnames:
        ws = formulas[name]
        cached = values[name]
        print(f"[SHEET] {name} size={ws.max_row}x{ws.max_column}")
        for row in ws.iter_rows():
            nonempty = [cell for cell in row if cell.value is not None]
            if not nonempty:
                continue
            labels = []
            for cell in nonempty[:12]:
                value = cell.value
                cached_value = cached[cell.coordinate].value
                if isinstance(value, str) and value.startswith("="):
                    labels.append(f"{cell.coordinate}={value} [cached={cached_value}]")
                else:
                    labels.append(f"{cell.coordinate}={clean(value)}")
            print(" | ".join(labels))
